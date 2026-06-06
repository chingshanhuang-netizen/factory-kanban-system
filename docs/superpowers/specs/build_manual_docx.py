"""
產生 TPS.Nexus 工廠看板系統操作手冊 .docx
輸出：2026-06-07-kanban-operation-manual.docx（同目錄）
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(__file__), "2026-06-07-kanban-operation-manual.docx")

# ── 顏色常數 ──────────────────────────────────────────────────────
C_DARK  = RGBColor(0x0d, 0x1b, 0x2a)
C_BLUE  = RGBColor(0x1a, 0x3a, 0x5c)
C_LIGHT = RGBColor(0x4f, 0xc3, 0xf7)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_TEXT  = RGBColor(0x1a, 0x1a, 0x2e)
C_HEAD  = RGBColor(0x0d, 0x47, 0xa1)
C_SUB   = RGBColor(0x15, 0x65, 0xc0)
C_NOTE  = RGBColor(0x37, 0x47, 0x4f)
C_GREEN = RGBColor(0x1b, 0x5e, 0x20)
C_TEAL  = RGBColor(0x00, 0x60, 0x64)
C_TH_BG = RGBColor(0x1a, 0x3a, 0x5c)

doc = Document()

# ── 頁面設定 ──────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = section.right_margin = Cm(2.5)
section.top_margin  = section.bottom_margin = Cm(2)

# ── 預設字型 ──────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = '微軟正黑體'
style.font.size = Pt(10.5)
style.font.color.rgb = C_TEXT
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

# ── 輔助函式 ──────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), val.get('sz', '4'))
            el.set(qn('w:color'), val.get('color', '1A3A5C'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_heading(doc, text, level, color=None):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.name = '微軟正黑體'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    if color:
        run.font.color.rgb = color
    if level == 0:
        run.font.size = Pt(20)
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = color or C_HEAD
    elif level == 2:
        run.font.size = Pt(12.5)
        run.font.bold = True
        run.font.color.rgb = color or C_SUB
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = color or C_TEAL
    return p

def add_para(doc, text, indent=0, bold=False, color=None, size=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.name = '微軟正黑體'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(doc, text, indent=1, sub=False):
    p = doc.add_paragraph(style='List Bullet' if not sub else 'List Bullet 2')
    p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.name = '微軟正黑體'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'bottom', 'left', 'right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), '1565C0')
        pBdr.append(el)
    pPr.append(pBdr)
    run = p.add_run('📌 ' + text)
    run.font.name = '微軟正黑體'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    run.font.color.rgb = RGBColor(0x0d, 0x47, 0xa1)
    run.font.size = Pt(10)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 標題列
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_bg(cell, '1A3A5C')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = C_WHITE
        run.font.name = '微軟正黑體'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
        run.font.size = Pt(10)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 資料列
    for r, row in enumerate(rows):
        bg = 'EBF5FB' if r % 2 == 0 else 'FFFFFF'
        for c, val in enumerate(row):
            cell = table.cell(r + 1, c)
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = '微軟正黑體'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
            run.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 欄寬
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'bottom', 'left', 'right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), 'AAAAAA')
        pBdr.append(el)
    pPr.append(pBdr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return p

def hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:color'), '1A3A5C')
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)

# ═══════════════════════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('⚙ TPS.Nexus 工廠看板系統')
run.font.name = '微軟正黑體'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = C_HEAD

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('作業操作手冊')
run2.font.name = '微軟正黑體'
run2._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
run2.font.size = Pt(26)
run2.font.bold = True
run2.font.color.rgb = C_DARK

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('版本 v1.3　　日期 2026-06-07\n適用對象：操作員、管理者、系統管理員')
run3.font.name = '微軟正黑體'
run3._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
run3.font.size = Pt(11)
run3.font.color.rgb = C_NOTE

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. 系統概述
# ═══════════════════════════════════════════════════════════════
add_heading(doc, '1. 系統概述', 1)

add_para(doc, 'TPS.Nexus 工廠看板系統讓操作員與管理者能在廠區樓層圖上即時掌握設備運行狀態。')
doc.add_paragraph()

add_heading(doc, '主要功能', 2)
for item in [
    '廠區平面圖上顯示各設備即時狀態（運行 / 待機 / 停機 / 警報）',
    '多地圖自動輪播（巡視模式，支援倒數徽章與暫停控制）',
    '點擊設備查看詳細參數、工單與警報紀錄',
    '管理者可拖曳設備位置、存草稿、發布佈局版本',
    '異常狀態即時警報推送（SignalR WebSocket）',
    '全螢幕顯示模式',
]:
    add_bullet(doc, item)

doc.add_paragraph()
add_heading(doc, '系統路由', 2)
add_table(doc,
    ['路由', '說明'],
    [
        ['/kanban/1、/kanban/2…', '各廠區看板地圖'],
        ['/kanban/settings', '設定管理（需 KANBAN_SETTINGS 權限）'],
    ],
    col_widths=[7, 9]
)

hr(doc)

# ═══════════════════════════════════════════════════════════════
# 2. 介面總覽
# ═══════════════════════════════════════════════════════════════
add_heading(doc, '2. 介面總覽', 1)

add_heading(doc, '頂部導覽列', 2)
add_table(doc,
    ['元素', '說明'],
    [
        ['⚙ TPS Nexus 看板系統', '系統名稱'],
        ['DEMO', '環境標示'],
        ['🗺 看板地圖（白色）', '目前在看板頁；任何地圖 URL 下均保持白色（active 狀態）'],
        ['⚙ 設定管理', '前往設定頁（需 KANBAN_SETTINGS 權限）'],
    ],
    col_widths=[5, 11]
)

add_heading(doc, '工具列按鈕', 2)
add_table(doc,
    ['按鈕', '說明', '所需權限'],
    [
        ['地圖名稱下拉', '切換至其他廠區地圖', 'KANBAN_VIEW'],
        ['✏️ 編輯', '進入 / 離開編輯模式', 'KANBAN_EDIT'],
        ['💾 存草稿', '儲存目前佈局為草稿（編輯模式中顯示）', 'KANBAN_EDIT'],
        ['🚀 發布', '將草稿發布為正式版本（編輯模式中顯示）', 'KANBAN_PUBLISH'],
        ['📋 版本', '開啟 / 關閉版本歷史面板', 'KANBAN_VIEW'],
        ['⛶ 全螢幕', '切換全螢幕顯示', 'KANBAN_VIEW'],
    ],
    col_widths=[4, 8, 4]
)

hr(doc)

# ═══════════════════════════════════════════════════════════════
# 3. 看板地圖 — 基本操作
# ═══════════════════════════════════════════════════════════════
add_heading(doc, '3. 看板地圖 — 基本操作', 1)

add_heading(doc, '3.1 切換地圖', 2)
add_heading(doc, '方式一：工具列下拉選單', 3)
for s in ['點擊工具列左側的地圖名稱下拉', '從清單中選擇目標廠區', '頁面自動切換，URL 更新為 /kanban/{新地圖 ID}']:
    add_bullet(doc, s)

add_heading(doc, '方式二：頂部「🗺 看板地圖」連結', 3)
add_bullet(doc, '點擊後跳回上次使用的地圖（透過 localStorage 記憶）')
add_note(doc, '系統會記住您上次查看的地圖，下次進入時自動跳轉。')

doc.add_paragraph()
add_heading(doc, '3.2 設備狀態燈號', 2)
add_table(doc,
    ['燈號顏色', '意義'],
    [
        ['綠色（發光）', '設備運行中'],
        ['橘色（發光）', '待機'],
        ['灰色', '停機'],
        ['紅色（閃爍）', '警報中（Widget 整體背景轉深紅，右上角顯示 ! 角標）'],
    ],
    col_widths=[4, 12]
)

hr(doc)

# ═══════════════════════════════════════════════════════════════
# 4. 地圖輪播
# ═══════════════════════════════════════════════════════════════
add_heading(doc, '4. 地圖輪播', 1)
add_para(doc, '輪播功能讓多廠區地圖自動循環顯示，適用於大螢幕巡視場景。')

add_heading(doc, '4.1 倒數徽章說明', 2)
add_table(doc,
    ['顯示', '意義'],
    [
        ['⏱ N 秒後切換', '輪播進行中，N 秒後跳到下一張地圖'],
        ['⏸ 輪播已暫停 + ▶ 繼續', '輪播已暫停'],
    ],
    col_widths=[6, 10]
)

add_heading(doc, '4.2 暫停 / 繼續', 2)
add_bullet(doc, '手動暫停：點擊徽章右側「⏸ 暫停」按鈕')
add_bullet(doc, '繼續輪播：點擊「▶ 繼續」，從剩餘秒數繼續倒數')

add_heading(doc, '4.3 自動暫停場景', 2)
add_table(doc,
    ['場景', '原因'],
    [
        ['開啟版本面板（📋 版本）', '避免版本操作中途切換地圖'],
        ['進入編輯模式', '編輯時停止輪播'],
    ],
    col_widths=[7, 9]
)

add_heading(doc, '4.4 自動恢復場景', 2)
add_table(doc,
    ['場景', '恢復說明'],
    [
        ['關閉版本面板', '從暫停時的剩餘秒數繼續'],
        ['完成編輯（發布或按「✅ 完成編輯」）', '重新啟動輪播'],
        ['地圖切換（輪播觸發的導覽）', '以新地圖的設定秒數重啟'],
    ],
    col_widths=[7, 9]
)

add_note(doc, '切換至新地圖時，暫停狀態會自動重置為「未暫停」。')
hr(doc)

# ═══════════════════════════════════════════════════════════════
# 5. 全螢幕模式
# ═══════════════════════════════════════════════════════════════
add_heading(doc, '5. 全螢幕模式', 1)

add_heading(doc, '5.1 進入全螢幕', 2)
add_para(doc, '點擊工具列右側「⛶」按鈕，地圖區域擴展至整個螢幕。')

add_heading(doc, '5.2 全螢幕下的功能', 2)
for item in [
    '✅ 輪播倒數徽章（右下角可見）',
    '✅ 暫停 / 繼續輪播按鈕',
    '✅ 設備 Widget 點擊',
    '✅ 警報角標',
]:
    add_bullet(doc, item)

add_heading(doc, '5.3 離開全螢幕', 2)
add_bullet(doc, '按鍵盤 Esc，或再次點擊「⛶」按鈕')
hr(doc)

# ═══════════════════════════════════════════════════════════════
# 6. 設備資訊查看
# ═══════════════════════════════════════════════════════════════
add_heading(doc, '6. 設備資訊查看', 1)

add_heading(doc, '6.1 Hover 快覽（滑鼠停留）', 2)
add_para(doc, '將滑鼠移至設備 Widget 上方，自動顯示參數摘要浮窗，包含：')
for item in ['所有參數名稱與當前數值（最多 6 個）', '最近趨勢折線圖（小型）', '「點擊查看完整詳情 →」提示']:
    add_bullet(doc, item)

add_heading(doc, '6.2 點擊詳情抽屜', 2)
add_para(doc, '操作：點擊設備 Widget，右側滑出詳情抽屜。')
add_table(doc,
    ['頁籤', '內容'],
    [
        ['當前數值', '所有參數值列表 + 趨勢折線圖'],
        ['工單', '相關工單清單（可含連結）'],
        ['警報歷史', '歷史警報紀錄'],
        ['文件', '設備相關文件連結'],
        ['自訂連結', '內嵌 iframe 或外部連結按鈕'],
    ],
    col_widths=[4, 12]
)
add_note(doc, '頁籤由管理者在「設定管理」中透過「設備連結設定」配置，各設備可有不同頁籤組合。')
add_para(doc, '關閉抽屜：點擊抽屜以外的地圖區域。')
hr(doc)

# ═══════════════════════════════════════════════════════════════
# 7. 編輯模式
# ═══════════════════════════════════════════════════════════════
add_heading(doc, '7. 編輯模式', 1)
add_note(doc, '所需權限：KANBAN_EDIT（編輯）、KANBAN_PUBLISH（發布）')

add_heading(doc, '7.1 進入編輯模式', 2)
add_bullet(doc, '點擊工具列「✏️ 編輯」按鈕')
add_para(doc, '若存在未發布草稿，系統會彈出對話框：')
add_table(doc,
    ['選項', '行為'],
    [
        ['繼續編輯', '載入草稿的佈局繼續修改'],
        ['重新開始', '以目前發布版本為基礎建立全新草稿'],
        ['關閉對話框（X）', '取消，不進入編輯模式'],
    ],
    col_widths=[4, 12]
)
add_para(doc, '進入後：輪播自動停止，設備 Widget 邊框變為虛線（可拖曳）。')

add_heading(doc, '7.2 拖曳設備', 2)
add_bullet(doc, '在編輯模式中，直接拖曳設備 Widget 至新位置')
add_bullet(doc, '放開後，位置即時儲存至草稿')
add_note(doc, '系統自動儲存每次拖曳的位置，不需要手動儲存每一步。')

add_heading(doc, '7.3 儲存草稿', 2)
add_para(doc, '點擊工具列「💾 存草稿」：')
add_bullet(doc, '將目前佈局儲存為草稿版本')
add_bullet(doc, '草稿不影響其他使用者看到的正式版本')
add_bullet(doc, '可多次儲存草稿（版本號遞增）')

add_heading(doc, '7.4 發布', 2)
add_para(doc, '點擊工具列「🚀 發布」：')
add_bullet(doc, '將目前草稿設為正式發布版本')
add_bullet(doc, '舊的正式版本自動封存（Archived）')
add_bullet(doc, '所有使用者立即看到新佈局')
add_bullet(doc, '離開編輯模式，輪播自動重啟')

add_heading(doc, '7.5 離開編輯模式（不發布）', 2)
add_para(doc, '點擊「✅ 完成編輯」：')
add_bullet(doc, '退出編輯模式')
add_bullet(doc, '草稿保留，下次進入編輯模式時可選擇繼續')
add_bullet(doc, '輪播自動重啟')
hr(doc)

# ═══════════════════════════════════════════════════════════════
# 8. 版本管理
# ═══════════════════════════════════════════════════════════════
add_heading(doc, '8. 版本管理', 1)

add_heading(doc, '8.1 開啟版本面板', 2)
add_para(doc, '點擊工具列「📋 版本」：')
add_bullet(doc, '面板從右側疊加於地圖上（地圖不縮小）')
add_bullet(doc, '輪播自動暫停')
add_bullet(doc, '顯示目前地圖的所有佈局版本')

add_heading(doc, '8.2 版本狀態說明', 2)
add_table(doc,
    ['狀態標籤', '顏色', '說明'],
    [
        ['已發布', '綠色', '目前正式版本（同時只有一個）'],
        ['封存', '灰色', '歷史版本，可回溯'],
        ['草稿', '藍色', '未發布的進行中版本'],
    ],
    col_widths=[4, 3, 9]
)

add_heading(doc, '8.3 回溯版本', 2)
for s in [
    '在版本面板找到目標封存版本',
    '點擊該版本右側「回溯」按鈕',
    '系統將該版本設為新的「已發布」版本',
    '原「已發布」版本自動封存',
]:
    add_bullet(doc, s)
add_note(doc, '回溯後，所有使用者立即看到回溯後的佈局。可再次回溯至任意封存版本以撤銷操作。')

add_heading(doc, '8.4 關閉版本面板', 2)
add_para(doc, '再次點擊工具列「📋 版本」，輪播自動恢復。')
hr(doc)

# ═══════════════════════════════════════════════════════════════
# 9. 即時警報
# ═══════════════════════════════════════════════════════════════
add_heading(doc, '9. 即時警報', 1)

add_heading(doc, '9.1 警報通知', 2)
add_para(doc, '警報條件觸發時，畫面右下角自動彈出通知（持續 5 秒後消失）。')
add_table(doc,
    ['警報等級', '通知顏色', 'Widget 狀態'],
    [
        ['Critical（緊急）', '紅色', '背景深紅 + 角標 + 紅燈閃爍'],
        ['Warning（警告）', '橘色', '背景深橘 + 角標 + 紅燈閃爍'],
        ['Info（資訊）', '藍色', '角標'],
    ],
    col_widths=[4, 4, 8]
)

add_heading(doc, '9.2 多人共享', 2)
add_para(doc, '系統使用 SignalR 即時推送，所有已連線的使用者同步收到相同警報通知，不需要重新整理頁面。')
hr(doc)

# ═══════════════════════════════════════════════════════════════
# 10. 設定管理
# ═══════════════════════════════════════════════════════════════
add_heading(doc, '10. 設定管理', 1)
add_note(doc, '所需權限：KANBAN_SETTINGS。前往路由 /kanban/settings 或點擊頂部「⚙ 設定管理」。')

add_heading(doc, '10.1 設備管理', 2)
add_heading(doc, '新增設備', 3)
add_bullet(doc, '點擊「+ 新增設備」')
add_table(doc,
    ['欄位', '說明', '必填'],
    [
        ['名稱', '設備顯示名稱（如「CNC-01」）', '✅'],
        ['標籤', '設備代號或識別碼', ''],
        ['說明', '設備描述', ''],
        ['圖示類型', 'CssClass（Font Awesome）或 CustomImage', '✅'],
        ['圖示值', 'CSS class 名稱或圖片路徑', ''],
    ],
    col_widths=[4, 10, 2]
)
add_heading(doc, '編輯設備', 3)
add_bullet(doc, '在列表中點擊「編輯」，修改後點擊「儲存」')
add_heading(doc, '刪除設備', 3)
add_bullet(doc, '點擊「刪除」，確認後從系統移除')
add_note(doc, '刪除設備前，請確認已從所有地圖佈局中移除，否則可能造成顯示異常。')

add_heading(doc, '10.2 資料來源', 2)
add_para(doc, '設定設備參數的資料來源，支援四種類型：')
add_table(doc,
    ['類型', '說明'],
    [
        ['SQL', '連接 MySQL 資料庫，強制使用參數化查詢'],
        ['CSV', '讀取 CSV 檔案第一筆資料列為最新數值'],
        ['JSON', '讀取 JSON 檔案，支援路徑導覽'],
        ['XML', '讀取 XML 檔案，支援 XPath'],
    ],
    col_widths=[3, 13]
)
add_heading(doc, 'SQL 資料來源範例', 3)
add_code(doc, '類型: SQL\n查詢: SELECT speed, temp FROM machine WHERE id = @machineId\n參數: {"machineId": 1}')
add_note(doc, 'SQL 查詢強制使用參數化查詢，請勿在查詢語句中直接拼入變數。')

add_heading(doc, '10.3 圖示上傳', 2)
add_para(doc, '上傳自訂設備圖示，供「CustomImage」類型使用。')
add_table(doc,
    ['項目', '說明'],
    [
        ['支援格式', 'PNG、JPG、JPEG、SVG'],
        ['檔案大小上限', '5 MB'],
    ],
    col_widths=[4, 12]
)
add_para(doc, '操作步驟：')
for s in ['點擊「選擇檔案」', '選取圖示檔案', '點擊「上傳」', '複製回傳的圖片路徑', '填入設備管理的「圖示值」欄位']:
    add_bullet(doc, s)

add_heading(doc, '10.4 地圖管理', 2)
add_heading(doc, '匯入新地圖', 3)
add_table(doc,
    ['格式', '說明'],
    [
        ['PNG 圖片', '廠區平面圖截圖'],
        ['JPG 圖片', '同上'],
        ['SVG 向量圖', '可縮放向量圖，線條清晰'],
        ['DXF 工程圖', 'AutoCAD 格式，自動轉換為 SVG'],
        ['JSON 座標定義', '純數據定義設備位置，無底圖'],
        ['XML 座標定義', '同上，XML 格式'],
    ],
    col_widths=[4, 12]
)
add_heading(doc, '已匯入地圖列表欄位', 3)
add_table(doc,
    ['欄位', '說明'],
    [
        ['名稱', '地圖名稱（取自上傳的檔案名）'],
        ['格式', 'PNG / JPG / SVG / DXF / JsonCoord / XmlCoord'],
        ['建立時間', '匯入日期'],
        ['版本', '管理版本標示（如 v1、v2）'],
    ],
    col_widths=[4, 12]
)
hr(doc)

# ═══════════════════════════════════════════════════════════════
# 11. 權限說明
# ═══════════════════════════════════════════════════════════════
add_heading(doc, '11. 權限說明', 1)
add_table(doc,
    ['權限碼', '可執行操作', '適用角色'],
    [
        ['KANBAN_VIEW', '查看看板地圖、設備狀態、版本面板（唯讀）', '所有使用者'],
        ['KANBAN_EDIT', '進入編輯模式、拖曳設備、存草稿', '管理者'],
        ['KANBAN_PUBLISH', '發布草稿、回溯版本', '資深管理者'],
        ['KANBAN_SETTINGS', '設備管理、資料來源、圖示上傳、地圖管理', '系統管理員'],
    ],
    col_widths=[4, 9, 3]
)
hr(doc)

# ═══════════════════════════════════════════════════════════════
# 12. 常見問題
# ═══════════════════════════════════════════════════════════════
add_heading(doc, '12. 常見問題', 1)

faqs = [
    ('Q1：切換地圖後，看板地圖按鈕變成藍色？',
     '正常行為（v1.3 已修正）。「🗺 看板地圖」連結在任何 /kanban/{n} 頁面下均保持白色。若顯示藍色，請重新整理頁面。'),
    ('Q2：全螢幕模式下，輪播倒數計時消失了？',
     '舊版本問題，v1.3 已修正。倒數徽章現在位於地圖容器內部，全螢幕模式下正常顯示。請清除瀏覽器快取並重新整理。'),
    ('Q3：版本面板關掉後輪播沒有繼續？',
     '確認：① 開啟面板 → 輪播暫停（正常）② 再次點擊「📋 版本」關閉 → 輪播自動恢復。若仍未恢復，確認是否同時進入了編輯模式（編輯模式下輪播停止）。'),
    ('Q4：我拖曳了設備但忘了發布，其他人看到的是舊版嗎？',
     '是的。草稿不影響正式版本，只有點擊「🚀 發布」後所有使用者才會看到新佈局。'),
    ('Q5：我誤發布了錯誤的佈局，如何還原？',
     '使用版本回溯功能：開啟「📋 版本」→ 找到正確的封存版本 → 點擊「回溯」，立即生效。'),
    ('Q6：警報通知消失後，在哪裡查看歷史警報？',
     '點擊任意設備 Widget，在詳情抽屜中選擇「警報歷史」頁籤。'),
    ('Q7：地圖下拉選單的地圖名稱顯示不正確？',
     '地圖名稱取自匯入時的檔案名稱。如需修改，請前往「設定管理 → 地圖管理」刪除後重新匯入。'),
    ('Q8：系統連線中斷後，警報還會推送嗎？',
     'SignalR 連線中斷時警報推送暫停，並顯示黃色警告通知「警報連線失敗」。請重新整理頁面重建連線。'),
]

for q, a in faqs:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.font.bold = True
    run.font.color.rgb = C_HEAD
    run.font.name = '微軟正黑體'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    add_para(doc, a, indent=0.5)
    doc.add_paragraph()

# ── 頁尾 ──────────────────────────────────────────────────────
hr(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('本手冊對應系統版本 v1.3（2026-06-07）　如有問題，請聯繫系統管理員。')
run.font.name = '微軟正黑體'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
run.font.size = Pt(9)
run.font.color.rgb = C_NOTE

doc.save(OUT)
print('OK: ' + OUT)
